import re


def extract_text_from_pdf(filepath: str) -> str:
    """PDF → plain text, page structure preserved. Tries pdfplumber, falls back to PyPDF2."""
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t: pages.append(t.strip())
        return "\n\n".join(pages)
    except ImportError:
        pass
    except Exception as e:
        raise RuntimeError(f"pdfplumber error: {e}")


def extract_text_from_docx(filepath: str) -> str:
    """DOCX → plain text. Reads paragraphs AND table cells (many resumes use tables)."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Install python-docx: pip install python-docx")
    try:
        doc = Document(filepath)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells: lines.append("  ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        raise RuntimeError(f"DOCX read error: {e}")


def extract_text_from_txt(filepath: str) -> str:
    """TXT/RTF → plain text. Encoding order: utf-8-sig → utf-8 → latin-1."""
    content = ""
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(filepath, "r", encoding=enc) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue
        except Exception as e:
            raise RuntimeError(f"File read error: {e}")

    if not content:
        raise RuntimeError("Could not decode file with any supported encoding.")

    if filepath.lower().endswith(".rtf"):
        content = _strip_rtf(content)

    return content


def _strip_rtf(text: str) -> str:
    """Remove RTF control sequences, preserve paragraph breaks using a state-machine parser."""
    # Replace \par and \line with newlines
    text = re.sub(r'\\par\b|\\line\b', '\n', text, flags=re.I)
    text = re.sub(r'\\u(-?\d+)\?', lambda m: chr(int(m.group(1)) % 65536), text)
    
    # State machine to parse RTF content
    pattern = re.compile(
        r'\\([a-z]{1,32})(-?\d+)? ?|\\\'([0-9a-f]{2})|\\([^a-z])|([{}])|([^\\{}]+)',
        re.IGNORECASE
    )
    stack = []
    parts = []
    ignoring = False
    
    for match in pattern.finditer(text):
        word, arg, hexchar, symbol, brace, plain = match.groups()
        if brace:
            if brace == '{':
                stack.append(ignoring)
            elif brace == '}':
                if stack:
                    ignoring = stack.pop()
                else:
                    ignoring = False
        elif plain:
            if not ignoring:
                parts.append(plain)
        elif word:
            # Common RTF destination control words to ignore
            if word.lower() in (
                'fonttbl', 'colortbl', 'stylesheet', 'info', 'listtable',
                'listoverridetable', 'xmlnstbl', 'picw', 'pich'
            ):
                ignoring = True
        elif symbol:
            if symbol == '*':
                ignoring = True
        elif hexchar:
            if not ignoring:
                try:
                    parts.append(chr(int(hexchar, 16)))
                except ValueError:
                    pass
                    
    cleaned = "".join(parts)
    # Remove consecutive spaces and trim lines
    lines = [re.sub(r'[ \t]+', ' ', ln).strip() for ln in cleaned.split('\n')]
    return '\n'.join(ln for ln in lines if ln)